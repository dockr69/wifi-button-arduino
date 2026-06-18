#!/usr/bin/env python3
"""Erzeugt ein handschriftlich ausfüllbares A4-Formular (Funkauslöser)
im P.O.S. Corporate Design 2023 — mit echter Kopf-/Fußzeile."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

N_TASTER = 10

# ── Corporate Design ─────────────────────────────────────────────────────────
DUNKELBLAU = RGBColor(0x1B, 0x3B, 0x62)   # Überschriften
POS_BLAU = RGBColor(0x09, 0x50, 0x9D)     # Akzent / Striche
DUNKELGRAU = RGBColor(0x45, 0x57, 0x64)   # Fließtext / Sublines
HELLGRAU = "EBEBEB"                        # Flächen / Tabellenkopf
ZEBRA = "F5F7FA"                           # dezenter Zeilenwechsel
FONT = "Outfit"            # Regular (Fließtext)
F_BLACK = "Outfit Black"   # Überschriften (CD)
F_SEMI = "Outfit SemiBold" # Sublines / Labels (CD)
F_MED = "Outfit Medium"
DL = "/Users/fabi/Downloads"
HEADER_IMG = f"{DL}/240409_POS_Kopfzeile_NEU_transparent.png"
FOOTER_IMG = f"{DL}/240402_POS_Fußzeile_NEU_transparent.png"

PAGE_W = Cm(21.0)


def fixed_table_layout(table, widths):
    """Erzwingt feste Spaltenbreiten (sonst verteilt Word gleichmäßig)."""
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # Gesamtbreite fixieren
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(sum(w.emu / 914400 * 1440 for w in widths))))
    tblPr.append(tblW)
    # tblGrid neu setzen
    for el in tbl.findall(qn("w:tblGrid")):
        tbl.remove(el)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.emu / 914400 * 1440)))
        grid.append(gc)
    tbl.insert(0, grid)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def accent_rule(p, width_pt=2, color="09509D"):
    """Kurzer P.O.S.-blauer Strich unter den Absatz (CD: Überschrift + Strich)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(width_pt * 8)))
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def set_letter_spacing(run, pt):
    """Laufweite (Tracking) in pt — für edlere Versal-Überschriften."""
    rPr = run._r.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(int(pt * 20)))  # twips
    rPr.append(sp)


def full_bleed_img(paragraph, path, neg_indent):
    """Bild randlos über die volle Seitenbreite platzieren."""
    paragraph.paragraph_format.left_indent = -neg_indent
    no_space(paragraph)
    paragraph.add_run().add_picture(path, width=PAGE_W)


doc = Document()

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(1.6)
sec.top_margin = Cm(3.6)        # Platz unter der Kopfzeile
sec.bottom_margin = Cm(3.2)     # Platz über der Fußzeile
sec.header_distance = Cm(0.0)
sec.footer_distance = Cm(0.0)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10)
style.font.color.rgb = DUNKELGRAU

# ── Kopf- / Fußzeile (randlos, volle Breite) ─────────────────────────────────
if os.path.exists(HEADER_IMG):
    full_bleed_img(sec.header.paragraphs[0], HEADER_IMG, sec.left_margin)
if os.path.exists(FOOTER_IMG):
    full_bleed_img(sec.footer.paragraphs[0], FOOTER_IMG, sec.left_margin)


# ── Helpers ──────────────────────────────────────────────────────────────────
def section_header(text):
    hd = doc.add_paragraph(); no_space(hd)
    hd.paragraph_format.space_before = Pt(10)
    hd.paragraph_format.space_after = Pt(6)
    rr = hd.add_run(text.upper())
    rr.font.name = F_SEMI; rr.font.size = Pt(11)
    rr.font.color.rgb = DUNKELBLAU
    set_letter_spacing(rr, 0.6)
    accent_rule(hd, width_pt=2)
    return hd


def label(p, text):
    r = p.add_run(text); r.font.name = F_SEMI
    r.font.color.rgb = DUNKELBLAU; return r


def fill(p, text):
    r = p.add_run(text); r.font.color.rgb = DUNKELGRAU; return r


def hint(p, text):
    r = p.add_run(text); r.font.size = Pt(8); r.font.color.rgb = DUNKELGRAU
    return r


# ── Titel ────────────────────────────────────────────────────────────────────
h = doc.add_paragraph(); no_space(h)
h.paragraph_format.right_indent = Cm(11)   # kurzer Akzentstrich statt voller Breite
r = h.add_run("Funkauslöser")
r.font.name = F_BLACK; r.font.size = Pt(28)
r.font.color.rgb = DUNKELBLAU
accent_rule(h, width_pt=2.5)

sub = doc.add_paragraph(); no_space(sub)
sub.paragraph_format.space_before = Pt(5)
rs = sub.add_run("Konfigurations-Formular")
rs.font.name = F_SEMI; rs.font.size = Pt(11)
rs.font.color.rgb = DUNKELGRAU
doc.add_paragraph().add_run().font.size = Pt(4)

# ── Kunde / Datum ────────────────────────────────────────────────────────────
p = doc.add_paragraph(); no_space(p); p.paragraph_format.space_after = Pt(8)
label(p, "Kunde / Standort: "); fill(p, "_" * 45 + "      ")
label(p, "Datum: "); fill(p, "_" * 16)

# ── WLAN ─────────────────────────────────────────────────────────────────────
section_header("WLAN")

p = doc.add_paragraph(); no_space(p); p.paragraph_format.space_after = Pt(8)
label(p, "WLAN-Name (SSID): "); fill(p, "_" * 34 + "      ")
label(p, "Passwort: "); fill(p, "_" * 24)

p = doc.add_paragraph(); no_space(p); p.paragraph_format.space_after = Pt(8)
label(p, "IP-Vergabe: ")
rr = p.add_run("☐ DHCP (automatisch)      ☐ Statisch")
rr.font.size = Pt(11); rr.font.color.rgb = DUNKELBLAU
hint(p, "   → bei „Statisch\" unten je Taster eine IP eintragen")

p = doc.add_paragraph(); no_space(p); p.paragraph_format.space_after = Pt(8)
label(p, "Gateway: "); fill(p, "_" * 26 + "      ")
label(p, "Subnetz: "); fill(p, "_" * 26)
hint(p, "   (nur bei statischer IP)")

# ── Taster-Tabelle ───────────────────────────────────────────────────────────
section_header("Taster")

headers = ["Nr.", "Bezeichnung / Textname", "IP-Adresse", "Aktion / Ansage"]
widths = [Cm(0.9), Cm(6.1), Cm(4.0), Cm(6.8)]

table = doc.add_table(rows=1 + N_TASTER, cols=len(headers))
table.style = "Table Grid"
fixed_table_layout(table, widths)

hdr = table.rows[0].cells
for i, (txt, w) in enumerate(zip(headers, widths)):
    hdr[i].width = w
    set_cell_bg(hdr[i], "1B3B62")  # dunkelblauer Kopf, weiße Schrift
    para = hdr[i].paragraphs[0]; no_space(para)
    run = para.add_run(txt)
    run.font.name = F_SEMI; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for n in range(1, N_TASTER + 1):
    cells = table.rows[n].cells
    for i, w in enumerate(widths):
        cells[i].width = w
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if n % 2 == 0:
            set_cell_bg(cells[i], ZEBRA)
        para = cells[i].paragraphs[0]; no_space(para)
        if i == 0:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(n)); run.font.name = F_SEMI
            run.font.color.rgb = DUNKELBLAU

for n in range(1, N_TASTER + 1):
    table.rows[n].height = Cm(1.05)
    table.rows[n].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

# ── Fußnote ──────────────────────────────────────────────────────────────────
f = doc.add_paragraph(); no_space(f); f.paragraph_format.space_before = Pt(8)
hint(f, "Hinweis: „Aktion / Ansage\" = was beim Druck passieren soll "
        "(z. B. Name der Durchsage / MP3-Datei).")

out = "/Users/fabi/development/wifi-button-arduino/wifi-button-builder/Funkausloeser_Formular.docx"
doc.save(out)
print("gespeichert:", out)
